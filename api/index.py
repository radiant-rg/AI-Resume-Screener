
import os
import json
import tempfile
from pathlib import Path

from flask import Flask, request, jsonify, render_template
from flask_cors import CORS

from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field
from pypdf import PdfReader
from docx import Document


# ============================================================
# 1. FLASK SETUP
# ============================================================



BASE_DIR = Path(__file__).resolve().parent.parent

app = Flask(
    __name__,
    template_folder=str(BASE_DIR / "templates"),
    static_folder=str(BASE_DIR / "static")
)

CORS(app)
# Allows frontend to communicate with Flask
CORS(app)


# ============================================================
# 2. LOAD GROQ API
# ============================================================

load_dotenv()

my_api_key = os.getenv("GROQ_API_KEY")
print("API key loaded:", bool(my_api_key))

if not my_api_key:
    raise ValueError("GROQ_API_KEY is missing.")

client = Groq(api_key=my_api_key)

model = "minimaxai/minimax-m2.7"


# ============================================================
# 3. JOB DESCRIPTION MODEL
# ============================================================

class JobD(BaseModel):
    role: str
    required_skills: list[str] = Field(default_factory=list)
    preferred_skills: list[str] = Field(default_factory=list)
    minimum_experience: float | None = None
    education_requirements: list[str] = Field(default_factory=list)
    responsibilities: list[str] = Field(default_factory=list)


jobd_schema = JobD.model_json_schema()


# ============================================================
# 4. EXPERIENCE MODEL
# ============================================================

class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = Field(default_factory=list)


# ============================================================
# 5. RESUME MODEL
# ============================================================

class Resume(BaseModel):

    name: str | None = None

    email: str | None = None

    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = Field(default_factory=list)

    experiences: list[Experience] = Field(
        default_factory=list
    )

    education: list[str] = Field(
        default_factory=list
    )

    projects: list[str] = Field(
        default_factory=list
    )

    certifications: list[str] = Field(
        default_factory=list
    )


resume_schema = Resume.model_json_schema()


# ============================================================
# 6. MATCH RESULT
# ============================================================

class MatchResult(BaseModel):

    score: float

    details: dict


# ============================================================
# 7. READ PDF
# ============================================================

def read_pdf(file_path):

    reader = PdfReader(file_path)

    text = ""

    for page in reader.pages:

        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text


# ============================================================
# 8. READ DOCX
# ============================================================

def read_docx(file_path):

    document = Document(file_path)

    text = ""

    # Paragraphs
    for paragraph in document.paragraphs:

        if paragraph.text.strip():

            text += paragraph.text + "\n"

    # Tables
    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                if cell.text.strip():

                    text += cell.text + "\n"

    return text


# ============================================================
# 9. READ RESUME
# ============================================================

def read_resume(file_path):

    extension = file_path.suffix.lower()

    if extension == ".pdf":

        return read_pdf(file_path)

    elif extension == ".docx":

        return read_docx(file_path)

    else:

        return None


# ============================================================
# 10. PARSE JOB DESCRIPTION
# ============================================================

def parse_job_description(job_description):

    system_prompt = f"""
You are an expert HR assistant.

Analyze the job description and extract structured information.

Return ONLY valid JSON matching this schema:

{json.dumps(jobd_schema, indent=2)}

Rules:

1. Do not invent information.

2. Extract the actual job role.

3. Extract required skills.

4. Extract preferred skills.

5. Extract minimum experience if explicitly mentioned.

6. Extract education requirements.

7. Extract key responsibilities.

8. If information is missing, use null for optional values
   and an empty list for list fields.
"""

    user_prompt = f"""
Analyze this job description:

{job_description}
"""

    response = client.chat.completions.create(

        model=model,

        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_prompt
            }
        ],

        response_format={
            "type": "json_object"
        }
    )

    raw_output = response.choices[0].message.content

    data = json.loads(raw_output)

    return JobD(**data)


# ============================================================
# 11. PARSE RESUME
# ============================================================

def parse_resume(resume_text):

    system_prompt = f"""
You are an expert resume parser.

Extract information from the resume based on its meaning.

Return ONLY valid JSON matching this schema:

{json.dumps(resume_schema, indent=2)}

Rules:

1. Do not invent information.

2. If information is unavailable, return null.

3. Empty lists should be returned as [].

4. Include internships inside experiences.

5. Extract skills from the entire resume.

6. Extract projects.

7. Extract certifications.

8. Extract education.

9. Extract candidate contact information when available.
"""

    user_prompt = f"""
Parse this resume:

{resume_text}
"""

    response = client.chat.completions.create(

        model=model,

        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_prompt
            }
        ],

        response_format={
            "type": "json_object"
        }
    )

    raw_output = response.choices[0].message.content

    data = json.loads(raw_output)

    return Resume(**data)


# ============================================================
# 12. COMPARE RESUME WITH JOB
# ============================================================

def final_score(job, resume):

    match_schema = MatchResult.model_json_schema()

    prompt = f"""
You are an expert technical recruiter.

Compare the candidate resume against the job description.

JOB DESCRIPTION:

{job.model_dump_json(indent=2)}


CANDIDATE:

{resume.model_dump_json(indent=2)}


Return JSON matching this schema:

{json.dumps(match_schema, indent=2)}


The score must be between 0 and 100.

The details object MUST contain:

- candidate_name
- matching_skills
- missing_important_skills
- experience_requirement_met
- overall_match_percentage
- final_verdict

Evaluate realistically.

Give more importance to required skills than
preferred skills.

Consider projects and internships as relevant
experience where appropriate.

Do not invent information.
"""

    response = client.chat.completions.create(

        model=model,

        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],

        response_format={
            "type": "json_object"
        }
    )

    data = json.loads(
        response.choices[0].message.content
    )

    return MatchResult(**data)


# ============================================================
# 13. HEALTH CHECK
# ============================================================

@app.route("/" , methods=["GET"])
def home():
    return render_template("index.html")


# ============================================================
# 14. SCREEN RESUMES
# ============================================================

@app.route("/screen", methods=["POST"])
def screen_resumes():

    try:

        # ----------------------------------------------------
        # GET JOB DESCRIPTION
        # ----------------------------------------------------

        job_description = request.form.get(
            "job_description"
        )

        if not job_description:

            return jsonify({
                "success": False,
                "error": "Job description is required."
            }), 400


        # ----------------------------------------------------
        # GET RESUME FILES
        # ----------------------------------------------------

        files = request.files.getlist(
            "resumes"
        )

        if not files:

            return jsonify({
                "success": False,
                "error": "Please upload at least one resume."
            }), 400


        # ----------------------------------------------------
        # PARSE JOB
        # ----------------------------------------------------

        print("\nAnalyzing job description...")

        job = parse_job_description(
            job_description
        )


        print("\nJOB ROLE:")
        print(job.role)

        print("\nREQUIRED SKILLS:")
        print(job.required_skills)


        # ----------------------------------------------------
        # PROCESS RESUMES
        # ----------------------------------------------------

        all_results = []


        for file in files:

            print(
                f"\nProcessing: {file.filename}"
            )


            # ------------------------------------------------
            # Check extension
            # ------------------------------------------------

            filename = file.filename

            extension = Path(filename).suffix.lower()

            if extension not in [".pdf", ".docx"]:

                print(
                    "Skipping unsupported file:",
                    filename
                )

                continue


            # ------------------------------------------------
            # Create temporary file
            # ------------------------------------------------

            with tempfile.NamedTemporaryFile(
                delete=False,
                suffix=extension
            ) as temp_file:

                file.save(
                    temp_file.name
                )

                temp_path = Path(
                    temp_file.name
                )


            try:

                # --------------------------------------------
                # Extract text
                # --------------------------------------------

                resume_text = read_resume(
                    temp_path
                )


                if not resume_text:

                    print(
                        "Could not read:",
                        filename
                    )

                    continue


                # --------------------------------------------
                # Parse resume
                # --------------------------------------------

                parsed_resume = parse_resume(
                    resume_text
                )


                print(
                    "Candidate:",
                    parsed_resume.name
                )


                # --------------------------------------------
                # Compare with job
                # --------------------------------------------

                result = final_score(
                    job,
                    parsed_resume
                )


                print(
                    "Score:",
                    result.score
                )


                # --------------------------------------------
                # Store result
                # --------------------------------------------

                all_results.append({

                    "name": parsed_resume.name,

                    "email": parsed_resume.email,

                    "phone": parsed_resume.phone,

                    "score": result.score,

                    "skills": parsed_resume.skills,

                    "education": parsed_resume.education,

                    "projects": parsed_resume.projects,

                    "certifications": parsed_resume.certifications,

                    "experiences": [
                        experience.model_dump()
                        for experience
                        in parsed_resume.experiences
                    ],

                    "details": result.details,

                    "resume_file": filename
                })


            finally:

                # Delete temporary file
                if temp_path.exists():

                    temp_path.unlink()


        # ----------------------------------------------------
        # SORT RESULTS
        # ----------------------------------------------------

        all_results.sort(
            key=lambda candidate:
            candidate["score"],
            reverse=True
        )


        # ----------------------------------------------------
        # RETURN EVERYTHING TO FRONTEND
        # ----------------------------------------------------

        return jsonify({

            "success": True,

            "job": {

                "role": job.role,

                "required_skills":
                    job.required_skills,

                "preferred_skills":
                    job.preferred_skills,

                "minimum_experience":
                    job.minimum_experience,

                "education_requirements":
                    job.education_requirements,

                "responsibilities":
                    job.responsibilities
            },

            "total_candidates":
                len(all_results),

            "candidates":
                all_results,

            "top_candidates":
                all_results[:2],

            "lowest_candidates":
                all_results[-2:]
        })


    except Exception as e:

        print(
            "ERROR:",
            str(e)
        )

        return jsonify({

            "success": False,

            "error": str(e)
        }), 500


# ============================================================
# 15. START SERVER
# ============================================================

if __name__ == "__main__":

    app.run(debug=True)